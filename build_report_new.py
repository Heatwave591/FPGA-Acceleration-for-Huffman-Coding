import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from PIL import Image

# ── Helpers ───────────────────────────────────────────────────────────────────

def add_before(doc, ref_para, text='', bold=False, italic=False, center=False, space_before=0, space_after=8):
    """Add paragraph at doc end (properly connected), move before ref_para."""
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    ref_para._p.addprevious(p._p)
    return p

def add_image_before(doc, ref_para, path, caption):
    img = Image.open(path)
    w_px, h_px = img.size
    dpi = 96
    w_in = w_px / dpi
    h_in = h_px / dpi
    scale = min(5.5 / w_in, 3.2 / h_in, 1.0)
    w_in *= scale

    # Image
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.space_after = Pt(2)
    img_p.add_run().add_picture(path, width=Inches(w_in))
    ref_para._p.addprevious(img_p._p)

    # Caption
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_after = Pt(10)
    run = cap_p.add_run(caption)
    run.italic = True
    ref_para._p.addprevious(cap_p._p)

def add_table_before(doc, ref_para, title, headers, rows):
    """Insert a proper Word table before ref_para."""
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(8)
    title_p.paragraph_format.space_after  = Pt(4)
    r = title_p.add_run(title)
    r.bold = True
    ref_para._p.addprevious(title_p._p)

    # Table
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = 'Table Grid'

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = tbl.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = cell_text
            row_cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Move table before ref_para
    ref_para._p.addprevious(tbl._tbl)

    # Spacer after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    ref_para._p.addprevious(spacer._p)

# ── Open base ─────────────────────────────────────────────────────────────────
doc = Document('Report_clean.docx')
print('Opened Report_clean.docx -', len(doc.paragraphs), 'paragraphs')

# Locate anchors
p_iv_impl     = None
p_correctness = None
p_discussion  = None
p_references  = None

for p in doc.paragraphs:
    t = p.text.strip()
    if t == 'IV. IMPLEMENTATION'     and not p_iv_impl:     p_iv_impl     = p
    if t == 'D. Correctness Summary' and not p_correctness: p_correctness = p
    if 'DISCUSSION' in t             and not p_discussion:  p_discussion  = p
    if t == 'REFERENCES'             and not p_references:  p_references  = p

assert all([p_iv_impl, p_correctness, p_discussion, p_references]), 'Anchor not found!'
print('All anchors found.')

# ── 1. Section III.E — Operational Workflow ───────────────────────────────────
items_iiie = [
    ('heading', 'E. Operational Workflow', {}),
    ('body', 'Fig. 6 illustrates the encoder workflow. The encoder is first configured by writing a '
             '(length, code) pair into each of the 256 code_table entries. Once configuration is '
             'complete, on each cycle where symbol_valid and symbol_ready are both asserted, the '
             'encoder looks up the Huffman code and shifts it MSB-first into a 64-bit accumulator. '
             'When the accumulator holds at least 32 bits, the upper 32 bits are emitted as '
             'stream_out and the accumulator is left-shifted by 32. At end of stream, additional '
             'zero-valued symbols flush any residual bits.', {}),
    ('center_italic', '[Fig. 6 - Encoder operational flowchart]', {}),
    ('body', 'Fig. 7 illustrates the decoder workflow. The decoder is configured in three phases: '
             'first_code table, base table, and symbol_table. During decoding, each incoming 32-bit '
             'stream word is loaded into the top of a 64-bit sliding window. The decoder finds the '
             'smallest length L where the top-16-bit peek >= first_code[L], reads the symbol from '
             'symbol_table[base[L] + ((peek - first_code[L]) >> (16-L))], and advances the window '
             'by L bits.', {}),
    ('center_italic', '[Fig. 7 - Decoder operational flowchart]', {}),
    ('blank', '', {}),
]

for kind, text, opts in items_iiie:
    if kind == 'heading':
        add_before(doc, p_iv_impl, text, bold=True, italic=True, space_before=14, space_after=6)
    elif kind == 'body':
        add_before(doc, p_iv_impl, text, space_after=8)
    elif kind == 'center_italic':
        add_before(doc, p_iv_impl, text, italic=True, center=True, space_after=8)
    elif kind == 'blank':
        add_before(doc, p_iv_impl, '', space_after=4)

print('Added III.E.')

# ── 2. Section V.D — Gradient Distribution Simulation Results ─────────────────
waveforms = [
    ('screenshots/Highly_Sparse/Encoder_config_Phase.png',
     'Fig. 8. Highly sparse - encoder configuration phase (0-200 ns). enc_data_config = 10000 confirms 1-bit code for 0x00.'),
    ('screenshots/Highly_Sparse/Encoding_and_first_flush.png',
     'Fig. 9. Highly sparse - encoding and decoder output (5400-5900 ns). Avg code length 3.3 bits, 2.41x compression.'),
    ('screenshots/Highly_Sparse/Decoder_output.png',
     'Fig. 10. Highly sparse - decoder output (5650-5950 ns). symbol_out correctly reproduces input sequence.'),
    ('screenshots/Moderate_sparse/Encoder config phase.png',
     'Fig. 11. Moderate - encoder configuration phase (0-200 ns). enc_data_config = 30000 confirms 3-bit code for 0x00.'),
    ('screenshots/Moderate_sparse/Encoding and first flush.png',
     'Fig. 12. Moderate - encoding and decoder output (5400-5900 ns). Avg code length 5.9 bits, 1.35x compression.'),
    ('screenshots/Moderate_sparse/Decoder_output.png',
     'Fig. 13. Moderate - decoder output (5600-5900 ns). Round-trip verified.'),
    ('screenshots/Dense/Encoder config phase.png',
     'Fig. 14. Dense - encoder configuration phase (0-200 ns). enc_data_config = 50000 confirms 5-bit code for 0x00.'),
    ('screenshots/Dense/Encoding and first flush.png',
     'Fig. 15. Dense - encoding and decoder output (5400-5900 ns). Avg code length 7.9 bits, 1.02x compression.'),
    ('screenshots/Dense/Decoder output.png',
     'Fig. 16. Dense - decoder output (5650-5950 ns). Round-trip verified.'),
]

add_before(doc, p_correctness, 'D. Gradient Distribution Simulation Results',
           bold=True, italic=True, space_before=14, space_after=6)
add_before(doc, p_correctness,
    'The following waveforms demonstrate hardware correctness across all three gradient '
    'distributions. For each scenario: the encoder configuration phase (0-200 ns) confirms '
    'the correct codebook was loaded; the 5400-5900 ns window shows encoding and round-trip '
    'behaviour; and the close-up confirms symbol_out matches the original input.',
    space_after=8)

add_before(doc, p_correctness, '1) Highly Sparse  (FREQS[0x00] = 10,000  |  Compression Ratio: 2.41x)',
           bold=True, space_before=10, space_after=4)
for path, cap in waveforms[0:3]:
    add_image_before(doc, p_correctness, path, cap)

add_before(doc, p_correctness, '2) Moderate  (FREQS[0x00] = 1,000  |  Compression Ratio: 1.35x)',
           bold=True, space_before=10, space_after=4)
for path, cap in waveforms[3:6]:
    add_image_before(doc, p_correctness, path, cap)

add_before(doc, p_correctness, '3) Dense  (FREQS[0x00] = 100  |  Compression Ratio: 1.02x)',
           bold=True, space_before=10, space_after=4)
for path, cap in waveforms[6:9]:
    add_image_before(doc, p_correctness, path, cap)

add_before(doc, p_correctness, '', space_after=4)
print('Added V.D with 9 waveforms.')

# ── 3. Section VI — Gradient Compression Performance Analysis ─────────────────
add_before(doc, p_discussion, '', space_after=4)

add_table_before(doc, p_discussion,
    'TABLE II - GRADIENT TRANSFER IMPACT (GPT-2, 248 MB, 1 Gbps LINK)',
    ['Scenario', 'Compressed Size', 'Transfer Time', 'Time Saved', 'Reduction'],
    [
        ['Highly Sparse (2.41x)', '103 MB', '823 ms',  '1161 ms', '58.5%'],
        ['Moderate (1.35x)',      '184 MB', '1470 ms', '514 ms',  '25.9%'],
        ['Dense (1.02x)',         '243 MB', '1945 ms', '39 ms',   '2.0%' ],
        ['No compression',        '248 MB', '1984 ms', '-',       '-'    ],
    ]
)

add_before(doc, p_discussion,
    'Under sparse gradient conditions - the dominant case during early training - FPGA '
    'Huffman compression reduces the 248 MB payload to approximately 103 MB, cutting '
    'transfer time by 58.5% per synchronisation round. For moderate gradients the saving '
    'is 25.9%. For dense gradients the saving is negligible (2.0%), and bypassing '
    'compression is preferable to avoid encoding latency.', space_after=8)

add_before(doc, p_discussion,
    'For a GPT-2 model (124 million parameters, bfloat16 precision), each gradient '
    'synchronisation round transmits approximately 248 MB. Table II shows the compressed '
    'payload and estimated transfer time on a 1 Gbps gRPC link (baseline: 1984 ms).', space_after=8)

add_before(doc, p_discussion, 'C. Impact on Gradient Communication',
           bold=True, italic=True, space_before=14, space_after=6)

add_table_before(doc, p_discussion,
    'TABLE I - HUFFMAN COMPRESSION ACROSS GRADIENT SCENARIOS',
    ['Scenario', 'FREQS[0x00]', 'Avg Code Length (bits)', 'Compression Ratio'],
    [
        ['Highly Sparse', '10,000', '3.314', '2.41x'],
        ['Moderate',       '1,000', '5.908', '1.35x'],
        ['Dense',            '100', '7.862', '1.02x'],
    ]
)

add_before(doc, p_discussion,
    'Table I summarises the canonical Huffman results for each scenario. '
    'The average code length is the weighted mean over all 256 symbols; '
    'the compression ratio is defined as 8 / average_code_length.', space_after=8)

add_before(doc, p_discussion, 'B. Compression Results',
           bold=True, italic=True, space_before=14, space_after=6)

add_before(doc, p_discussion,
    'The highly sparse scenario (FREQS[0x00] = 10000) models early-training gradients '
    'where zero bytes dominate. The moderate scenario (FREQS[0x00] = 1000) models '
    'mid-training gradients. The dense scenario (FREQS[0x00] = 100) models late-training '
    'gradients where the byte distribution is nearly uniform.', space_after=8)
add_before(doc, p_discussion,
    'The 256-symbol Huffman hardware is designed to compress bfloat16 gradient tensors '
    'before gRPC transmission in a distributed LLM training system. Three frequency '
    'distributions were defined that model different stages of training, and a canonical '
    'Huffman codebook was computed for each using Python.', space_after=8)

add_before(doc, p_discussion, 'A. Motivation and Scenario Definition',
           bold=True, italic=True, space_before=14, space_after=6)

add_before(doc, p_discussion, 'VI. GRADIENT COMPRESSION PERFORMANCE ANALYSIS',
           bold=True, center=True, space_before=14, space_after=8)

print('Added VI. Gradient Compression Performance Analysis.')

# ── 4. GitHub note before REFERENCES ─────────────────────────────────────────
github_p = doc.add_paragraph()
github_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
github_p.paragraph_format.space_after = Pt(8)
github_p.add_run(
    'The complete source code for this project - including the SystemVerilog encoder '
    'and decoder, the system testbench, and the Python Huffman table generation scripts '
    '- is publicly available at: '
)
r = github_p.add_run('https://github.com/[your-username]/FPGA-Acceleration-for-Huffman-Coding')
r.bold = True
p_references._p.addprevious(github_p._p)
print('Added GitHub note.')

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save('report_new.docx')
print('\nDone. report_new.docx saved.')
