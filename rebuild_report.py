import sys, io, zipfile, re, os, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Step 1: Strip all drawings from unpacked3 doc and repack as base ──────────
with open('unpacked3/word/document.xml', encoding='utf-8') as f:
    doc_xml = f.read()

# Remove all <w:drawing>...</w:drawing> blocks
doc_xml_clean = re.sub(r'<w:drawing>.*?</w:drawing>', '', doc_xml, flags=re.DOTALL)
print('Drawings removed. Remaining count:', doc_xml_clean.count('w:drawing'))

with open('unpacked3/word/document.xml', 'w', encoding='utf-8') as f:
    f.write(doc_xml_clean)

# Repack in correct order
def repack(src_dir, out_file):
    all_files = []
    for root, dirs, files in os.walk(src_dir):
        for f in files:
            abs_path = os.path.join(root, f)
            arc_path = os.path.relpath(abs_path, src_dir).replace(os.sep, '/')
            all_files.append((arc_path, abs_path))

    def sort_key(item):
        arc = item[0]
        if arc == '[Content_Types].xml': return (0, arc)
        elif arc.startswith('_rels/'): return (1, arc)
        elif arc.startswith('word/_rels/'): return (2, arc)
        elif arc == 'word/document.xml': return (3, arc)
        elif arc.startswith('word/'): return (4, arc)
        else: return (5, arc)

    all_files.sort(key=sort_key)
    with zipfile.ZipFile(out_file, 'w', zipfile.ZIP_DEFLATED) as zf:
        for arc_path, abs_path in all_files:
            zf.write(abs_path, arc_path)

repack('unpacked3', 'Report_base.docx')
print('Repacked to Report_base.docx (text only, no images)')

# ── Step 2: Use python-docx to add waveform images ───────────────────────────
EMU_PER_INCH = 914400
MAX_W_IN = 5.5
MAX_H_IN = 3.2
DPI = 96

def get_dims(path):
    img = Image.open(path)
    w_px, h_px = img.size
    w_in = w_px / DPI
    h_in = h_px / DPI
    scale = min(MAX_W_IN / w_in, MAX_H_IN / h_in, 1.0)
    return w_in * scale, h_in * scale

images = [
    ('screenshots/Highly_Sparse/Encoder_config_Phase.png',
     'Fig. 8. Highly sparse - encoder configuration phase (0-200 ns). enc_data_config = 10000 confirms 1-bit code for 0x00.'),
    ('screenshots/Highly_Sparse/Encoding_and_first_flush.png',
     'Fig. 9. Highly sparse - encoding and decoder output (5400-5900 ns). stream_valid pulses infrequent; avg code length 3.3 bits (2.41x compression).'),
    ('screenshots/Highly_Sparse/Decoder_output.png',
     'Fig. 10. Highly sparse - close-up decoder output (5650-5950 ns). symbol_out correctly reproduces input sequence.'),
    ('screenshots/Moderate_sparse/Encoder config phase.png',
     'Fig. 11. Moderate - encoder configuration phase (0-200 ns). enc_data_config = 30000 confirms 3-bit code for 0x00.'),
    ('screenshots/Moderate_sparse/Encoding and first flush.png',
     'Fig. 12. Moderate - encoding and decoder output (5400-5900 ns). stream_valid more frequent; avg code length 5.9 bits (1.35x compression).'),
    ('screenshots/Moderate_sparse/Decoder_output.png',
     'Fig. 13. Moderate - close-up decoder output (5600-5900 ns). Round-trip verified for moderate distribution.'),
    ('screenshots/Dense/Encoder config phase.png',
     'Fig. 14. Dense - encoder configuration phase (0-200 ns). enc_data_config = 50000 confirms 5-bit code for 0x00.'),
    ('screenshots/Dense/Encoding and first flush.png',
     'Fig. 15. Dense - encoding and decoder output (5400-5900 ns). stream_valid most frequent; avg code length 7.9 bits (1.02x compression).'),
    ('screenshots/Dense/Decoder output.png',
     'Fig. 16. Dense - close-up decoder output (5650-5950 ns). Round-trip verified for dense distribution.'),
]

doc = Document('Report_base.docx')

# Find paragraph containing 'D. Correctness Summary' and insert before it
target_idx = None
for i, p in enumerate(doc.paragraphs):
    if 'D. Correctness Summary' in p.text:
        target_idx = i
        break

if target_idx is None:
    print('ERROR: Could not find D. Correctness Summary')
    exit(1)

print('Inserting images before paragraph', target_idx, ':', doc.paragraphs[target_idx].text[:50])

# We'll add images at the end of section V by appending after the last
# waveform paragraph in section V (before D. Correctness Summary)
# Use python-docx's add_picture on a paragraph inserted before the target

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def insert_paragraph_before(doc, ref_para, text='', bold=False, italic=False, center=False):
    new_p = OxmlElement('w:p')
    ref_para._p.addprevious(new_p)
    # Get the new paragraph object
    for i, p in enumerate(doc.paragraphs):
        if p._p is new_p:
            para = p
            break
    else:
        # Fallback
        import docx.text.paragraph
        para = docx.text.paragraph.Paragraph(new_p, doc)
    if text:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para

def insert_picture_before(doc, ref_para, img_path, width_in, height_in, caption):
    # Image paragraph
    img_p = insert_paragraph_before(doc, ref_para, center=True)
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_p.add_run()
    run.add_picture(img_path, width=Inches(width_in))

    # Caption paragraph
    cap_p = insert_paragraph_before(doc, ref_para, caption, italic=True, center=True)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return img_p, cap_p

ref_para = doc.paragraphs[target_idx]

# Insert section heading and intro (in reverse since each goes before ref_para)
# We'll build a list and insert in forward order using a moving reference

# First collect all paragraphs to insert as (type, args)
to_insert = []

# Section heading
to_insert.append(('heading', 'D. Gradient Distribution Simulation Results'))
to_insert.append(('body', 'The following waveforms demonstrate hardware correctness across all three gradient distributions. For each scenario, three timestamps are shown: the encoder configuration phase (0-200 ns) confirming the correct codebook was loaded, the encoding and decoding window (5400-5900 ns) showing stream output and round-trip behaviour, and a close-up of the decoder output confirming symbol_out matches the original input sequence.'))
to_insert.append(('subhead', '1) Highly Sparse  (FREQS[0x00] = 10,000  |  Compression Ratio: 2.41x)'))
for i in range(3):
    to_insert.append(('image', images[i]))
to_insert.append(('subhead', '2) Moderate  (FREQS[0x00] = 1,000  |  Compression Ratio: 1.35x)'))
for i in range(3, 6):
    to_insert.append(('image', images[i]))
to_insert.append(('subhead', '3) Dense  (FREQS[0x00] = 100  |  Compression Ratio: 1.02x)'))
for i in range(6, 9):
    to_insert.append(('image', images[i]))

# Insert in reverse order so they appear in correct order before ref_para
for kind, content in reversed(to_insert):
    if kind == 'image':
        path, caption = content
        w_in, h_in = get_dims(path)
        # Caption first (reversed, so it ends up after image)
        cap_p = insert_paragraph_before(doc, ref_para, caption, italic=True, center=True)
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Image
        img_p = insert_paragraph_before(doc, ref_para, center=True)
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_p.add_run()
        run.add_picture(path, width=Inches(w_in))
    elif kind == 'heading':
        p = insert_paragraph_before(doc, ref_para, content, bold=True, italic=True)
    elif kind == 'subhead':
        p = insert_paragraph_before(doc, ref_para, content, bold=True)
    elif kind == 'body':
        p = insert_paragraph_before(doc, ref_para, content)

doc.save('Report.docx')
print('Done. Report.docx saved with all 9 waveforms.')
