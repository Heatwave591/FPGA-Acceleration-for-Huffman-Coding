import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from PIL import Image

doc = Document('report_new.docx')
print('Opened report_new.docx -', len(doc.paragraphs), 'paragraphs,', len(doc.tables), 'tables')

# ── Step 1: Remove all the wrongly-ordered added content ─────────────────────
# Keep only paragraphs that belong to the original Report_clean content.
# Original ends just before 'D. Correctness Summary' additions.
# Strategy: find first paragraph of each added section and remove everything added.

# Markers for content to REMOVE (added by the broken script)
remove_markers = [
    'D. Gradient Distribution Simulation Results',
    '1) Highly Sparse',
    '2) Moderate',
    '3) Dense',
    'VI. GRADIENT COMPRESSION PERFORMANCE ANALYSIS',
    'A. Motivation and Scenario Definition',
    'B. Compression Results',
    'C. Impact on Gradient Communication',
    'D. Hardware Verification Across Scenarios',
    'TABLE I',
    'TABLE II',
    'Under sparse gradient conditions',
    'E. Operational Workflow',
    'Fig. 6 illustrates the encoder',
    'Fig. 7 illustrates the decoder',
    '[Fig. 6',
    '[Fig. 7',
    'The complete source code for this project',
]

# Collect paragraph XML elements to remove
paras_to_remove = []
for p in doc.paragraphs:
    t = p.text.strip()
    for marker in remove_markers:
        if t.startswith(marker) or marker in t:
            paras_to_remove.append(p._p)
            break

# Remove tables (all added tables - the 2 word tables I added)
tables_to_remove = [tbl._tbl for tbl in doc.tables]

# Also remove image paragraphs (paragraphs with drawings that are Fig 8-16)
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith('Fig.') and any(f'Fig. {n}' in t for n in range(8, 17)):
        paras_to_remove.append(p._p)

print(f'Removing {len(paras_to_remove)} paragraphs and {len(tables_to_remove)} tables')

for elem in paras_to_remove:
    try:
        elem.getparent().remove(elem)
    except:
        pass

for elem in tables_to_remove:
    try:
        elem.getparent().remove(elem)
    except:
        pass

# Also remove image-only paragraphs (no text, but have w:drawing)
from lxml import etree
body = doc.element.body
for p in list(body):
    if p.tag.endswith('}p'):
        has_drawing = p.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
        text = ''.join(t.text or '' for t in p.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))
        if has_drawing is not None and not text.strip():
            body.remove(p)

print('After cleanup:', len(doc.paragraphs), 'paragraphs,', len(doc.tables), 'tables')

# Verify anchors
for p in doc.paragraphs:
    t = p.text.strip()
    if t in ('IV. IMPLEMENTATION', 'D. Correctness Summary', 'REFERENCES') or 'DISCUSSION' in t:
        print(f'  Anchor: {t}')

# ── Helpers ───────────────────────────────────────────────────────────────────

def add_para(doc, ref_para, text='', bold=False, italic=False,
             center=False, space_before=0, space_after=8):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    ref_para._p.addprevious(p._p)
    return p

def add_img(doc, ref_para, path, caption):
    img = Image.open(path)
    w_px, h_px = img.size
    scale = min(5.5 / (w_px/96), 3.2 / (h_px/96), 1.0)
    w_in = (w_px/96) * scale
    ip = doc.add_paragraph()
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.paragraph_format.space_after = Pt(2)
    ip.add_run().add_picture(path, width=Inches(w_in))
    ref_para._p.addprevious(ip._p)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(10)
    cp.add_run(caption).italic = True
    ref_para._p.addprevious(cp._p)

def add_table(doc, ref_para, title, headers, rows):
    # Title
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(8)
    tp.paragraph_format.space_after  = Pt(4)
    tp.add_run(title).bold = True
    ref_para._p.addprevious(tp._p)
    # Table
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]
            c.text = val
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref_para._p.addprevious(tbl._tbl)
    # Spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)
    ref_para._p.addprevious(sp._p)

# ── Re-locate anchors ─────────────────────────────────────────────────────────
p_iv_impl     = next(p for p in doc.paragraphs if p.text.strip() == 'IV. IMPLEMENTATION')
p_correctness = next(p for p in doc.paragraphs if p.text.strip() == 'D. Correctness Summary')
p_discussion  = next(p for p in doc.paragraphs if 'DISCUSSION' in p.text)
p_references  = next(p for p in doc.paragraphs if p.text.strip() == 'REFERENCES')

# ── 1. Section III.E — Operational Workflow (FORWARD order) ───────────────────
add_para(doc, p_iv_impl, 'E. Operational Workflow',
         bold=True, italic=True, space_before=14, space_after=6)
add_para(doc, p_iv_impl,
    'Fig. 6 illustrates the encoder workflow. The encoder is first configured by writing a '
    '(length, code) pair into each of the 256 code_table entries. Once configuration is '
    'complete, on each cycle where symbol_valid and symbol_ready are both asserted, the encoder '
    'looks up the Huffman code and shifts it MSB-first into a 64-bit accumulator. When the '
    'accumulator holds at least 32 bits, the upper 32 bits are emitted as stream_out and the '
    'accumulator is left-shifted by 32. At end of stream, additional zero-valued symbols flush '
    'any residual bits.')
add_para(doc, p_iv_impl, '[Fig. 6 - Encoder operational flowchart]',
         italic=True, center=True, space_after=8)
add_para(doc, p_iv_impl,
    'Fig. 7 illustrates the decoder workflow. The decoder is configured in three phases: '
    'first_code table, base table, and symbol_table. During decoding, each incoming 32-bit '
    'stream word is loaded into the top of a 64-bit sliding window. The decoder finds the '
    'smallest L where peek >= first_code[L], reads symbol_table[base[L] + ((peek - '
    'first_code[L]) >> (16-L))], and advances the window by L bits.')
add_para(doc, p_iv_impl, '[Fig. 7 - Decoder operational flowchart]',
         italic=True, center=True, space_after=8)
add_para(doc, p_iv_impl, '', space_after=4)
print('Added III.E.')

# ── 2. Section V.D — Gradient Distribution Waveforms (FORWARD order) ──────────
waveforms = [
    ('screenshots/Highly_Sparse/Encoder_config_Phase.png',
     'Fig. 8. Highly sparse - encoder configuration phase (0-200 ns). enc_data_config=10000 confirms 1-bit code for 0x00.'),
    ('screenshots/Highly_Sparse/Encoding_and_first_flush.png',
     'Fig. 9. Highly sparse - encoding and decoder output (5400-5900 ns). Avg code length 3.3 bits, 2.41x compression.'),
    ('screenshots/Highly_Sparse/Decoder_output.png',
     'Fig. 10. Highly sparse - decoder output (5650-5950 ns). symbol_out correctly reproduces input sequence.'),
    ('screenshots/Moderate_sparse/Encoder config phase.png',
     'Fig. 11. Moderate - encoder configuration phase (0-200 ns). enc_data_config=30000 confirms 3-bit code for 0x00.'),
    ('screenshots/Moderate_sparse/Encoding and first flush.png',
     'Fig. 12. Moderate - encoding and decoder output (5400-5900 ns). Avg code length 5.9 bits, 1.35x compression.'),
    ('screenshots/Moderate_sparse/Decoder_output.png',
     'Fig. 13. Moderate - decoder output (5600-5900 ns). Round-trip verified.'),
    ('screenshots/Dense/Encoder config phase.png',
     'Fig. 14. Dense - encoder configuration phase (0-200 ns). enc_data_config=50000 confirms 5-bit code for 0x00.'),
    ('screenshots/Dense/Encoding and first flush.png',
     'Fig. 15. Dense - encoding and decoder output (5400-5900 ns). Avg code length 7.9 bits, 1.02x compression.'),
    ('screenshots/Dense/Decoder output.png',
     'Fig. 16. Dense - decoder output (5650-5950 ns). Round-trip verified.'),
]

add_para(doc, p_correctness, 'D. Gradient Distribution Simulation Results',
         bold=True, italic=True, space_before=14, space_after=6)
add_para(doc, p_correctness,
    'The following waveforms demonstrate hardware correctness across all three gradient '
    'distributions. For each scenario: the encoder configuration phase (0-200 ns) confirms '
    'the correct codebook was loaded; the 5400-5900 ns window shows encoding and round-trip '
    'behaviour; and the close-up confirms symbol_out matches the original input.')
add_para(doc, p_correctness, '1) Highly Sparse  (FREQS[0x00] = 10,000  |  Compression Ratio: 2.41x)',
         bold=True, space_before=10, space_after=4)
for path, cap in waveforms[0:3]:
    add_img(doc, p_correctness, path, cap)
add_para(doc, p_correctness, '2) Moderate  (FREQS[0x00] = 1,000  |  Compression Ratio: 1.35x)',
         bold=True, space_before=10, space_after=4)
for path, cap in waveforms[3:6]:
    add_img(doc, p_correctness, path, cap)
add_para(doc, p_correctness, '3) Dense  (FREQS[0x00] = 100  |  Compression Ratio: 1.02x)',
         bold=True, space_before=10, space_after=4)
for path, cap in waveforms[6:9]:
    add_img(doc, p_correctness, path, cap)
add_para(doc, p_correctness, '', space_after=4)
print('Added V.D with 9 waveforms.')

# ── 3. Section VI — Gradient Compression Performance Analysis (FORWARD order) ──
add_para(doc, p_discussion, 'VI. GRADIENT COMPRESSION PERFORMANCE ANALYSIS',
         bold=True, center=True, space_before=14, space_after=8)
add_para(doc, p_discussion, 'A. Motivation and Scenario Definition',
         bold=True, italic=True, space_before=14, space_after=6)
add_para(doc, p_discussion,
    'The 256-symbol Huffman hardware is designed to compress bfloat16 gradient tensors before '
    'gRPC transmission in a distributed LLM training system. Three frequency distributions were '
    'defined that model different stages of training, and a canonical Huffman codebook was '
    'computed for each using Python.')
add_para(doc, p_discussion,
    'The highly sparse scenario (FREQS[0x00] = 10000) models early-training gradients where '
    'zero bytes dominate. The moderate scenario (FREQS[0x00] = 1000) models mid-training '
    'gradients. The dense scenario (FREQS[0x00] = 100) models late-training gradients where '
    'the byte distribution is nearly uniform.')
add_para(doc, p_discussion, '', space_after=4)
add_para(doc, p_discussion, 'B. Compression Results',
         bold=True, italic=True, space_before=14, space_after=6)
add_para(doc, p_discussion,
    'Table I summarises the canonical Huffman results for each scenario. The average code '
    'length is the weighted mean over all 256 symbols; the compression ratio is defined as '
    '8 / average_code_length.')
add_table(doc, p_discussion,
    'TABLE I - HUFFMAN COMPRESSION ACROSS GRADIENT SCENARIOS',
    ['Scenario', 'FREQS[0x00]', 'Avg Code Length (bits)', 'Compression Ratio'],
    [
        ['Highly Sparse', '10,000', '3.314', '2.41x'],
        ['Moderate',       '1,000', '5.908', '1.35x'],
        ['Dense',            '100', '7.862', '1.02x'],
    ])
add_para(doc, p_discussion,
    'The highly sparse distribution achieves 2.41x compression, driven by the dominant '
    'zero byte being assigned a single-bit code. The moderate distribution achieves 1.35x. '
    'The dense distribution achieves only 1.02x, confirming Huffman coding provides '
    'negligible benefit when all symbols are approximately equally likely.')
add_para(doc, p_discussion, '', space_after=4)
add_para(doc, p_discussion, 'C. Impact on Gradient Communication',
         bold=True, italic=True, space_before=14, space_after=6)
add_para(doc, p_discussion,
    'For a GPT-2 model (124 million parameters, bfloat16 precision), each gradient '
    'synchronisation round transmits approximately 248 MB. Table II shows the compressed '
    'payload and estimated transfer time on a 1 Gbps gRPC link (baseline: 1984 ms).')
add_table(doc, p_discussion,
    'TABLE II - GRADIENT TRANSFER IMPACT (GPT-2, 248 MB, 1 Gbps LINK)',
    ['Scenario', 'Compressed Size', 'Transfer Time', 'Time Saved', 'Reduction'],
    [
        ['Highly Sparse (2.41x)', '103 MB', '823 ms',  '1161 ms', '58.5%'],
        ['Moderate (1.35x)',      '184 MB', '1470 ms', '514 ms',  '25.9%'],
        ['Dense (1.02x)',         '243 MB', '1945 ms', '39 ms',   '2.0%' ],
        ['No compression',        '248 MB', '1984 ms', '-',       '-'    ],
    ])
add_para(doc, p_discussion,
    'Under sparse gradient conditions - the dominant case during early training - FPGA '
    'Huffman compression reduces the 248 MB payload to 103 MB, cutting transfer time by '
    '58.5% per synchronisation round. For moderate gradients the saving is 25.9%. For '
    'dense gradients the saving is negligible (2.0%).')
add_para(doc, p_discussion, '', space_after=4)
add_para(doc, p_discussion, 'D. Hardware Verification Across Scenarios',
         bold=True, italic=True, space_before=14, space_after=6)
add_para(doc, p_discussion,
    'Each scenario was verified end-to-end in Vivado behavioural simulation. The Python '
    'scripts regenerate huffman_tables.svh for each distribution; the testbench calls '
    'config_encoder_full() and config_decoder_full() at runtime - no RTL modifications '
    'are required between scenarios.')
add_para(doc, p_discussion,
    'In all three cases, symbol_out correctly reproduced the input sequence. The waveforms '
    'show the expected increase in stream_valid pulse frequency as average code length grows.')
add_para(doc, p_discussion, '', space_after=4)
print('Added VI.')

# ── 4. GitHub note ────────────────────────────────────────────────────────────
gp = doc.add_paragraph()
gp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
gp.paragraph_format.space_after = Pt(8)
gp.add_run(
    'The complete source code for this project - including the SystemVerilog encoder and '
    'decoder, the system testbench, and the Python Huffman table generation scripts - is '
    'publicly available at: ')
r = gp.add_run('https://github.com/[your-username]/FPGA-Acceleration-for-Huffman-Coding')
r.bold = True
p_references._p.addprevious(gp._p)
print('Added GitHub note.')

doc.save('report_new2.docx')
print('\nDone. report_new2.docx saved.')
